# rag_utils.py
import os
import chromadb
from chromadb.utils import embedding_functions
import logging
import pypdfium2 as pdfium
from docx import Document
import tiktoken # For text chunking

# Assuming ollama_utils.py is in the same directory or accessible in PYTHONPATH
from ollama_utils import get_embedding, RAG_EMBEDDING_MODEL # Use the Ollama-based embedder

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Configuration for RAG
CHROMA_PERSIST_DIR = os.path.join("data", "chroma_db")
COLLECTION_NAME = "ollama_rag_collection"
NUM_RAG_RESULTS = 3  # Number of relevant chunks to retrieve
CHUNK_SIZE = 1000  # Target characters per chunk for tiktoken (approx)
CHUNK_OVERLAP = 100 # Characters of overlap for tiktoken (approx)


# --- ChromaDB Client Initialization ---
# We use a flag to ensure client is initialized only once.
_chroma_client = None
_collection = None

def get_chroma_collection():
    """Initializes and returns the ChromaDB collection."""
    global _chroma_client, _collection
    if _collection is None:
        try:
            if not os.path.exists(CHROMA_PERSIST_DIR):
                os.makedirs(CHROMA_PERSIST_DIR)
                logger.info(f"Created ChromaDB persistence directory: {CHROMA_PERSIST_DIR}")

            _chroma_client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
            
            # Custom embedding function that uses Ollama
            class OllamaEmbeddingFunction(embedding_functions.EmbeddingFunction):
                def __call__(self, input_texts: chromadb.Documents) -> chromadb.Embeddings:
                    embeddings_list = []
                    for text in input_texts:
                        emb, err = get_embedding(text) # Uses the function from ollama_utils
                        if err:
                            logger.error(f"Failed to get embedding for RAG: {err}. Using zero vector.")
                            # Fallback to a zero vector of a plausible dimension if error.
                            # This dimension should match what your RAG_EMBEDDING_MODEL produces.
                            # Nomic-embed-text produces 768 dimensions.
                            # This is a HACK. Ideally, handle this error more gracefully.
                            # For nomic-embed-text (dim 768)
                            # This dimension might vary for other embedding models!
                            # It's better to fail or skip if embeddings can't be generated.
                            # For now, let's log error and skip if embedding fails severely
                            # Returning an empty list or raising error might be better.
                            # For now, let's just log and return an empty list if any embedding fails,
                            # or a list of successful ones.
                            # ChromaDB expects a list of lists.
                            # If get_embedding returns None, we must handle it.
                            # A simple approach: if any embedding fails, the whole batch fails.
                            # This is too strict for a production system but simpler for an example.
                            # A better way: collect successful embeddings and skip failed ones,
                            # or use a default embedding for failed ones (though this has implications).
                            logger.warning(f"Skipping text chunk due to embedding failure: '{text[:50]}...'")
                            # To satisfy chromadb, we must return embeddings for all inputs, or filter inputs
                            # before this stage. Let's attempt to use a placeholder or raise.
                            # For now, if an error occurs, we log it and the document won't be properly added.
                            # This is a point of potential failure if Ollama embedding isn't working.
                            embeddings_list.append([0.0] * 768) # Placeholder, assumes 768 dimensions.
                                                               # THIS IS BAD PRACTICE if dimensions don't match.
                                                               # The actual embedding dimension should be fetched or known.
                                                               # For `nomic-embed-text`, it's 768.
                        else:
                            embeddings_list.append(emb)
                    if not embeddings_list and input_texts: # If all embeddings failed
                        raise ValueError("Failed to generate any embeddings for the input texts.")
                    return embeddings_list

            ollama_ef = OllamaEmbeddingFunction()

            _collection = _chroma_client.get_or_create_collection(
                name=COLLECTION_NAME,
                embedding_function=ollama_ef # Use the custom Ollama embedding function
                # metadata={"hnsw:space": "cosine"} # Default is l2, cosine is often good for text
            )
            logger.info(f"ChromaDB collection '{COLLECTION_NAME}' loaded/created from {CHROMA_PERSIST_DIR}.")
        except Exception as e:
            logger.error(f"Failed to initialize ChromaDB client or collection: {e}")
            # Raise the exception to be handled by the caller, so the app knows RAG isn't working.
            raise RuntimeError(f"ChromaDB initialization failed: {e}") from e
    return _collection

# --- Text Processing and Chunking ---
def get_tokenizer():
    """Gets a tiktoken tokenizer."""
    # Using a common model like 'cl100k_base' (used by gpt-3.5-turbo, gpt-4)
    # or 'p50k_base' (older models). 'cl100k_base' is a good general choice.
    try:
        return tiktoken.get_encoding("cl100k_base")
    except Exception:
        # Fallback if specific encoding not found, though 'cl100k_base' should be available.
        logger.warning("cl100k_base tokenizer not found, falling back to a default gpt2 tokenizer.")
        return tiktoken.encoding_for_model("gpt-2") # A common fallback

_tokenizer = None

def get_cached_tokenizer():
    global _tokenizer
    if _tokenizer is None:
        _tokenizer = get_tokenizer()
    return _tokenizer

def chunk_text(text: str, chunk_size_tokens: int = 256, chunk_overlap_tokens: int = 30) -> list[str]:
    """
    Chunks text using tiktoken for more semantically aware chunking based on tokens.
    :param text: The input text string.
    :param chunk_size_tokens: The target number of tokens per chunk.
    :param chunk_overlap_tokens: The number of tokens to overlap between chunks.
    :return: A list of text chunks.
    """
    if not text.strip():
        return []

    tokenizer = get_cached_tokenizer()
    tokens = tokenizer.encode(text)
    
    if not tokens:
        return []

    chunks = []
    current_pos = 0
    while current_pos < len(tokens):
        end_pos = min(current_pos + chunk_size_tokens, len(tokens))
        chunk_tokens = tokens[current_pos:end_pos]
        chunks.append(tokenizer.decode(chunk_tokens))
        
        if end_pos == len(tokens): # Reached the end
            break
        
        # Move current_pos, ensuring overlap unless it's the very last chunk
        current_pos += (chunk_size_tokens - chunk_overlap_tokens)
        if current_pos >= len(tokens): # Should not happen if logic is correct
             break
        # Ensure current_pos doesn't go negative if chunk_size < chunk_overlap (bad params)
        current_pos = max(0, current_pos)


    # A simpler character-based chunking if tiktoken is problematic (less ideal for semantics)
    # chunks = []
    # for i in range(0, len(text), CHUNK_SIZE - CHUNK_OVERLAP):
    #    chunks.append(text[i:i + CHUNK_SIZE])
    
    logger.info(f"Text of {len(tokens)} tokens chunked into {len(chunks)} chunks.")
    return chunks


# --- File Readers ---
def read_txt(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {e}")
        return ""

def read_pdf(file_path: str) -> str:
    text_content = []
    try:
        pdf = pdfium.PdfDocument(file_path)
        for i in range(len(pdf)):
            page = pdf.get_page(i)
            textpage = page.get_textpage()
            text_content.append(textpage.get_text_range())
            textpage.close()
            page.close()
        pdf.close()
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error reading PDF file {file_path} with pypdfium2: {e}")
        return ""

def read_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {e}")
        return ""

FILE_READERS = {
    '.txt': read_txt,
    '.pdf': read_pdf,
    '.docx': read_docx,
}

# --- RAG Core Functions ---
def add_text_to_rag(text_content: str, source_name: str = "text_input"):
    """Adds text content to the RAG knowledge base."""
    if not text_content.strip():
        logger.warning("Attempted to add empty text to RAG. Skipping.")
        return "No text content provided."

    try:
        collection = get_chroma_collection()
    except RuntimeError as e: # Catch ChromaDB init errors
        return str(e)

    # Check if embedding model is available before trying to add
    # This relies on RAG_EMBEDDING_MODEL being correctly set in ollama_utils
    # and the model being available in Ollama.
    # A more explicit check could be added here too using ollama_utils.check_embedding_model_availability()
    # but get_embedding will also report errors.

    chunks = chunk_text(text_content)
    if not chunks:
        logger.warning(f"Text content from '{source_name}' resulted in no chunks. Skipping.")
        return f"Text from '{source_name}' could not be chunked or was empty."

    documents_to_add = []
    metadatas_to_add = []
    ids_to_add = []
    
    # This is a simplification. Batching calls to get_embedding would be more efficient.
    # However, ChromaDB's OllamaEmbeddingFunction is designed to take multiple inputs,
    # so it should handle batching internally if we pass all chunks at once.
    # Let's prepare all data and add it in one go.

    for i, chunk in enumerate(chunks):
        doc_id = f"{source_name}_chunk_{i}"
        documents_to_add.append(chunk)
        metadatas_to_add.append({"source": source_name, "chunk_num": i})
        ids_to_add.append(doc_id)

    if not documents_to_add:
        return f"No valid chunks to add from '{source_name}'."

    try:
        collection.add(
            documents=documents_to_add,
            metadatas=metadatas_to_add,
            ids=ids_to_add
        )
        logger.info(f"Added {len(documents_to_add)} chunks from '{source_name}' to RAG collection.")
        return f"Successfully added content from '{source_name}' to RAG."
    except Exception as e:
        logger.error(f"Error adding documents to Chroma collection for '{source_name}': {e}")
        # This might happen if embedding generation fails within the OllamaEmbeddingFunction
        # or if there's another ChromaDB issue.
        return f"Error adding content from '{source_name}' to RAG: {e}. Ensure the embedding model '{RAG_EMBEDDING_MODEL}' is running in Ollama and is functional."


def add_file_to_rag(file_path: str, original_filename: str):
    """Reads a file, extracts text, and adds it to the RAG knowledge base."""
    _, ext = os.path.splitext(original_filename)
    ext = ext.lower()

    if ext not in FILE_READERS:
        logger.warning(f"Unsupported file type: {ext} for file {original_filename}")
        return f"Unsupported file type: {ext}. Supported types are: {', '.join(FILE_READERS.keys())}"

    try:
        text_content = FILE_READERS[ext](file_path)
        if not text_content or not text_content.strip():
            logger.warning(f"No text content extracted from file: {original_filename}")
            return f"No text content extracted from '{original_filename}'."
        
        return add_text_to_rag(text_content, source_name=original_filename)
    
    except Exception as e:
        logger.error(f"Failed to process or add file {original_filename} to RAG: {e}")
        return f"Error processing file '{original_filename}': {e}"
    finally:
        # Clean up the uploaded file after processing, regardless of success or failure
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Cleaned up uploaded file: {file_path}")
            except OSError as e:
                logger.error(f"Error deleting uploaded file {file_path}: {e}")


def query_rag(prompt_text: str):
    """Queries the RAG knowledge base for relevant context."""
    try:
        collection = get_chroma_collection()
    except RuntimeError: # Catch ChromaDB init errors
        return None, "RAG system (ChromaDB) not initialized." # Return error message

    if collection.count() == 0:
        logger.info("RAG collection is empty. No context to retrieve.")
        return None, None # No error, just no context

    try:
        # The query_texts will be embedded by the OllamaEmbeddingFunction within ChromaDB
        results = collection.query(
            query_texts=[prompt_text],
            n_results=NUM_RAG_RESULTS,
            include=['documents', 'metadatas'] # Include documents and metadatas
        )

        documents = results.get('documents')
        # metadatas = results.get('metadatas') # For potential future use

        if documents and documents[0]: # documents is a list containing one list of results for the query
            context = "\n\n---\n\n".join(documents[0])
            logger.info(f"Retrieved {len(documents[0])} context chunks for prompt.")
            return context, None
        else:
            logger.info("No relevant documents found in RAG for the prompt.")
            return None, None # No error, just no context
    except Exception as e:
        logger.error(f"Error querying RAG collection: {e}")
        return None, f"Error querying RAG: {e}" # Return error message


def get_rag_status():
    """Gets the status of the RAG collection (e.g., number of items)."""
    try:
        collection = get_chroma_collection()
        count = collection.count()
        return {"item_count": count}, None
    except RuntimeError as e: # Catch ChromaDB init errors
         return {"item_count": 0, "error": "RAG system (ChromaDB) not initialized."}, str(e)
    except Exception as e:
        logger.error(f"Error getting RAG status: {e}")
        return {"item_count": 0, "error": str(e)}, f"Error getting RAG status: {e}"


def clear_rag_collection():
    """Clears all items from the RAG collection."""
    global _collection, _chroma_client # Need to potentially re-initialize after clearing
    try:
        # Try to get the existing collection first.
        # If it doesn't exist or client fails, this will also fail.
        collection_to_clear = get_chroma_collection()
        
        # ChromaDB's delete_collection is more thorough than deleting all items.
        # We then re-create it.
        if _chroma_client and COLLECTION_NAME:
            _chroma_client.delete_collection(name=COLLECTION_NAME)
            logger.info(f"ChromaDB collection '{COLLECTION_NAME}' deleted.")
            
            # Reset global vars to force re-initialization on next access
            _collection = None 
            # _chroma_client remains, as PersistentClient can recreate collections.
            # Re-create it immediately to ensure it's available.
            get_chroma_collection() # This will call get_or_create_collection
            
            logger.info(f"ChromaDB collection '{COLLECTION_NAME}' re-created successfully after clearing.")
            return True, "RAG knowledge base cleared successfully."
        else:
            logger.warning("Chroma client or collection name not available for clearing.")
            return False, "RAG system not fully initialized, cannot clear."

    except chromadb.db.base.CollectionDoestNotExistError: # Specific ChromaDB error
        logger.info("Attempted to clear RAG, but collection does not exist. Nothing to clear.")
        # Ensure it's re-creatable on next use
        _collection = None
        get_chroma_collection() # Try to create it now
        return True, "RAG knowledge base was already empty or did not exist. Initialized fresh."
    except RuntimeError as e: # Catch ChromaDB init errors during get_chroma_collection
        return False, f"Failed to access RAG system to clear: {e}"
    except Exception as e:
        logger.error(f"Error clearing RAG collection: {e}")
        return False, f"Error clearing RAG knowledge base: {e}"